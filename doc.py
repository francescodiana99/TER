# ================================== FUNCTIONS   =====================================
def delete_paragraph(paragraph):
    """this function deletes a paragraph"""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_empty_paragraphs(doc):
    """The function removes all the empty paragraphs before processing the document"""
    for p in doc.paragraphs:
        if p.text == '':
            delete_paragraph(paragraph=p)
    return doc



def get_drugs(p):
    """"This function aims to extract all the drugs from each paragraph"""
    drug = []
    for idx, r in enumerate(p.runs):
        # print(p.runs[0].text)
        # print(r.text)
        # I do not get why r is different from p.runs[0]
        if p.runs[0].text == r.text:
            prev = r
            if r.font.highlight_color == 3 or r.font.highlight_color == 5 :
                drug.append(r.text)
        else:
            # here I check if the run is a part of a longer expression, checking if the color of the highlight is the same of the previous run
            if ((r.font.highlight_color == 3 and prev.font.highlight_color == 3) or (r.font.highlight_color == 5 and prev.font.highlight_color == 5)):
                # if the condition is true, I append the text because it is part of the same expression
                drug[-1] = drug[-1] + r.text
            elif r.font.highlight_color == 3 or r.font.highlight_color == 5 :
                drug.append(r.text)
                # check if the element is the last one
            elif r.text == " " and p.runs[idx] != p.runs[-1]:
                # in case it is a space between two highlighted elements we add it to the list
                if (p.runs[idx - 1].font.highlight_color == 3 and p.runs[idx + 1].font.highlight_color == 3) or (p.runs[idx - 1].font.highlight_color == 5 and p.runs[idx + 1].font.highlight_color == 5):
                    drug.append(r.text)
            prev = r
    # now if there is a space in the list, we merge the text
    for idx, e in enumerate(drug):
        if e == " ":
            drug[idx - 1] = drug[idx - 1] + e + drug[idx + 1]
            del drug[idx]
            del drug[idx]

    # we create the tuples
    """to avoid issues with possible duplicates in the same paragraph, the idea is to extract
    the position of an annotation and then to cut the string to that index, avoiding that find() gives as
    result always the first occurrence of an entity"""
    # temporary text
    temp = p.text
    # current index in the paragraph
    index = 0
    # now we create the tuples
    annotations = []
    for e in drug:
        start = temp.find(e)
        end = start + len(e)
        annotations.append((start + index, end + index, "DRUG"))
        # save the index of the substring in p
        index = index + end
        # cut the part of p before the entity found
        temp = p.text[index:]

    return annotations


def get_adr(p):
    """"This function aims at extract all the adversal drug reactions from each paragraph"""
    adr = []
    for idx, r in enumerate(p.runs):
        # print(p.runs[0].text)
        # print(r.text)
        # I do not get why r is different from p.runs[0]
        if p.runs[0].text == r.text:
            prev = r
            if r.font.highlight_color == 4:
                adr.append(r.text)
        else:
            # here I check if the run is a part of a longer expression, checking if the color of the highlight is the same of the previous run
            if (r.font.highlight_color == 4 and prev.font.highlight_color == 4):
                # if the condition is true, I append the text because it is part of the same expression
                adr[-1] = adr[-1] + r.text
            elif r.font.highlight_color == 4:
                adr.append(r.text)
                # check if the element is the last one
            elif r.text == " " and p.runs[idx] != p.runs[-1]:
                # in case it is a space between two highlighted elements we added it to the list
                if p.runs[idx - 1].font.highlight_color == 4 and p.runs[idx + 1].font.highlight_color == 4:
                    adr.append(r.text)
            prev = r
    # now if there is a space in the list, we merge the text
    for idx, e in enumerate(adr):
        if e == " ":
            adr[idx - 1] = adr[idx - 1] + e + adr[idx + 1]
            del adr[idx]
            del adr[idx]

    # now we create the tuples
#    annotations = []
#    for e in adr:
#        start = p.text.find(e)
#        end = start + len(e)
        #  ADR stands for "adverse drug reaction"
#        annotations.append((start, end, "ADR"))
#    return annotations

    # now we create the tuples
    """to avoid issues with possible duplicates in the same paragraph, the idea is to extract
    the position of an annotation and then to cut the string to that index, avoiding that find() gives as
    result always the first occurrence of an entity"""
    # temporary text
    temp = p.text
    # current index in the paragraph
    index = 0
    # now we create the tuples
    annotations = []
    for e in adr:
        start = temp.find(e)
        end = start + len(e)
        annotations.append((start + index, end + index, "ADR"))
        # save the index of the substring in p
        index = index + end
        # cut the part of p before the entity found
        temp = p.text[index:]

    return annotations


def get_sym(p):
    """"This function extracts all the symptoms or diagnosis not linked to drugs"""
    sym = []
    for idx, r in enumerate(p.runs):
        # print(p.runs[0].text)
        # print(r.text)
        # I do not get why r is different from p.runs[0]
        if p.runs[0].text == r.text:
            prev = r
            if r.font.highlight_color == 16:
                sym.append(r.text)
        else:
            # here I check if the run is a part of a longer expression, checking if the color of the highlight is the same of the previous run
            if (r.font.highlight_color == 16 and prev.font.highlight_color == 16):
                # if the condition is true, I append the text because it is part of the same expression
                sym[-1] = sym[-1] + r.text
            elif r.font.highlight_color == 16:
                sym.append(r.text)
                # check if the element is the last one
            elif r.text == " " and p.runs[idx] != p.runs[-1]:
                # in case it is a space between two highlighted elements we added it to the list
                if p.runs[idx - 1].font.highlight_color == 16 and p.runs[idx + 1].font.highlight_color == 16:
                    sym.append(r.text)
            prev = r
    # now if there is a space in the list, we merge the text
    for idx, e in enumerate(sym):
        if e == " ":
            sym[idx - 1] = sym[idx - 1] + e + sym[idx + 1]
            del sym[idx]
            del sym[idx]

    # now we create the tuples
    """to avoid issues with possible duplicates in the same paragraph, the idea is to extract
    the position of an annotation and then to cut the string to that index, avoiding that find() gives as
    result always the first occurrence of an entity"""
    # temporary text
    temp = p.text
    # current index in the paragraph
    index = 0
    # now we create the tuples
    annotations = []
    for e in sym:
        start = temp.find(e)
        end = start + len(e)
        annotations.append((start + index, end + index, "NLD"))
        # save the index of the substring in p
        index = index + end
        # cut the part of p before the entity found
        temp = p.text[index:]

    return annotations


def preprocessing(doc):
    """this function extracts all the annotations from the document"""

    training_data = []
    for p in doc.paragraphs:
        drugs = get_drugs(p=p)
        adr = get_adr(p=p)
        sym = get_sym(p=p)
        ann = drugs + adr + sym
        training_data.append((p.text, ann))

#    count = 0
#    for text, ann in training_data:
#        if ann == []:
#            count = count + 1
#    print (len(training_data))
#    print(count)
#    print(len(training_data) - count)

    training_data = [ (text,ann) for (text,ann) in training_data if ann != [] ]
    return training_data

#-------------TEST------------

if __name__ == "__main__":
    import docx
    from docx import Document
    document = docx.Document("./data/debug.docx")
    for p in document.paragraphs:
        drugs = get_drugs(p)
        adr = get_adr(p)
        sym = get_sym(p)
        print(drugs)
